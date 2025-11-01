from langchain_groq import ChatGroq
from langchain.agents import create_react_agent, AgentExecutor
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, AIMessage

from smartdoc_agent.core.tools import (
    analyze_document_structure,
    create_formatting_plan,
    apply_contextual_formatting,
    validate_formatting_result,
    apply_business_proposal_template_tool
)
from smartdoc_agent.config import get_groq_api_keys

import json
import traceback
import time
from langchain_core.agents import AgentAction, AgentFinish
from typing import Any
from groq import APIError

class DocumentFormattingAgent:
    def __init__(self, model_name="llama-3.3-70b-versatile", temperature=0, progress_callback=None, disable_llm=False):
        self.progress_callback = progress_callback
        self._emit_progress({"type": "lifecycle", "event": "agent_init_start"})

        self.model_name = model_name
        self.temperature = temperature
        self.disable_llm = disable_llm

        if not self.disable_llm:
            try:
                self.api_keys = get_groq_api_keys()
                self.current_key_index = 0
            except ValueError as e:
                self._emit_progress({"type": "error", "message": f"Fatal: {e}"})
                raise
            self._initialize_llm()
        else:
            self.api_keys = []
            self.current_key_index = 0
            self._emit_progress({"type": "info", "message": "LLM initialization disabled for frontend testing."})

        self.tools = [
            analyze_document_structure,
            create_formatting_plan,
            apply_contextual_formatting,
            validate_formatting_result,
            apply_business_proposal_template_tool
        ]
        self.tool_names = [tool.name for tool in self.tools]

        prompt_template = """
        You are an expert document formatting assistant. Your goal is to help users format their .docx documents.
        You have access to the following tools:

        {tools}

        To use a tool, use the following format:
        Thought: Do I need to use a tool? Yes
        Action: The action to take. Must be one of [{tool_names}]
        Action Input: The input to the action. If the tool has multiple arguments, this must be a JSON object formatted as a multi-line block. If the tool takes a single string argument (like a file path), provide just the string.
        Observation: The result of the action.

        When you have a response to say to the Human, or if you do not need to use a tool, you MUST use the format:
        Thought: Do I need to use a tool? No
        Final Answer: [your response here]

        A few important rules:
        1. When a tool provides output (Observation), and you need to use that output as input for a subsequent tool, you MUST use the exact, complete output from the Observation for the relevant parameter, UNLESS a special placeholder is specified below.
        2. Ensure your Action Input is correctly formatted for the tool.
        3. For the 'create_formatting_plan' tool:
           - Its 'document_analysis_json' parameter requires a JSON object (a dictionary).
           - Look at the Observation from the 'analyze_document_structure' tool. This observation is a large JSON string.
           - Extract ONLY the essential information: 'document_path' and 'summary' from this observation.
           - DO NOT include the full 'elements' array as it can be very large and cause errors.
           - In your Action Input, provide a JSON object where 'document_analysis_json' contains only the essential summary information.
           Example:
           Observation: {{"elements": [...very long array...], "document_path": "path/doc.docx", "summary": {{"total_elements": 10, "paragraph_count": 5, "heading_count": 2}}, ...}}
           Thought: I need to make a plan. I will extract only the essential summary information to avoid large JSON issues.
           Action: create_formatting_plan
           Action Input: {{"document_analysis_json": {{"document_path": "path/doc.docx", "summary": {{"total_elements": 10, "paragraph_count": 5, "heading_count": 2}}}}, "user_goal": "the user's original goal text"}}
        4. For the 'apply_contextual_formatting' tool:
           - It needs 'doc_path', 'formatting_plan_json', 'document_analysis_json', and 'output_doc_path'.
           - 'doc_path' is the original input document path.
           - For 'formatting_plan_json', you MUST use the literal string placeholder: "$CURRENT_FORMATTING_PLAN".
           - For 'document_analysis_json', you MUST use the literal string placeholder: "$FULL_ORIGINAL_ANALYSIS".
           - 'output_doc_path' is where the formatted document should be saved.
           Example Action Input: {{"doc_path": "doc.docx", "formatting_plan_json": "$CURRENT_FORMATTING_PLAN", "document_analysis_json": "$FULL_ORIGINAL_ANALYSIS", "output_doc_path": "doc_formatted.docx"}}
        5. For the 'validate_formatting_result' tool:
           - It requires these arguments: 'original_doc_analysis_json', 'modified_doc_analysis_json', 'formatting_plan_json', and 'user_goal'.
           - For 'original_doc_analysis_json', you MUST use the literal string placeholder: "$FULL_ORIGINAL_ANALYSIS".
           - For 'modified_doc_analysis_json', you MUST use the literal string placeholder: "$FULL_MODIFIED_ANALYSIS" (this refers to the analysis of the document *after* formatting has been applied by `apply_contextual_formatting` and the modified document re-analyzed by `analyze_document_structure`).
           - For 'formatting_plan_json', you MUST use the literal string placeholder: "$CURRENT_FORMATTING_PLAN".
           - 'user_goal' is the original user goal text provided at the start of the task.
           Example Action Input for 'validate_formatting_result':
           {{"original_doc_analysis_json": "$FULL_ORIGINAL_ANALYSIS", "modified_doc_analysis_json": "$FULL_MODIFIED_ANALYSIS", "formatting_plan_json": "$CURRENT_FORMATTING_PLAN", "user_goal": "The user's original request text here"}}
        6. For the 'apply_business_proposal_template_tool':
           - Use this when the user wants to convert unstructured content into a standard business proposal format.
           - It requires 'input_doc_path' (source document) and 'output_doc_path' (where to save the formatted proposal).
           - Optional 'template_path' can specify a custom template.
           - This tool will extract business content (company name, project title, executive summary, etc.) and apply it to a professional template.
           Example Action Input: {{"input_doc_path": "path/to/raw/content.docx", "output_doc_path": "path/to/formatted/proposal.docx"}}

        Begin!

        Previous conversation history:
        {chat_history}

        New input: {input}
        {agent_scratchpad}
        """

        self.prompt = ChatPromptTemplate.from_template(prompt_template).partial(
            tools="\n".join([f"{tool.name}: {tool.description}" for tool in self.tools]),
            tool_names=", ".join([tool.name for tool in self.tools]),
        )

        react_agent = create_react_agent(
            llm=self.llm,
            tools=self.tools,
            prompt=self.prompt
        )
        
        # Create a callback to monitor what's happening inside AgentExecutor
        from langchain.callbacks.base import BaseCallbackHandler
        
        class DebugCallbackHandler(BaseCallbackHandler):
            def on_agent_action(self, action, **kwargs):
                print(f"DEBUG CALLBACK: Agent Action - Tool: {action.tool}, Input: {str(action.tool_input)[:200]}...")
                
            def on_agent_finish(self, finish, **kwargs):
                print(f"DEBUG CALLBACK: Agent Finish - Output: {str(finish.return_values)[:200]}...")
                
            def on_tool_start(self, serialized, input_str, **kwargs):
                print(f"DEBUG CALLBACK: Tool Start - {serialized.get('name', 'Unknown')}")
                
            def on_tool_end(self, output, **kwargs):
                print(f"DEBUG CALLBACK: Tool End - Output: {str(output)[:200]}...")
                
            def on_llm_start(self, serialized, prompts, **kwargs):
                print(f"DEBUG CALLBACK: LLM Start - Model: {serialized.get('name', 'Unknown')}")
                if prompts:
                    print(f"DEBUG CALLBACK: LLM Prompt preview: {str(prompts[0])[:300]}...")
                
            def on_llm_end(self, response, **kwargs):
                print(f"DEBUG CALLBACK: LLM End - Response: {str(response.generations[0][0].text)[:200]}...")
        
        debug_callback = DebugCallbackHandler()
        
        # Set global reference for tools to access agent data
        import smartdoc_agent.core.tools as tools_module
        tools_module._current_agent_instance = self
        
        self.agent = AgentExecutor(
            agent=react_agent,
            tools=self.tools,
            verbose=True,
            max_iterations=15,
            handle_parsing_errors=True,
            callbacks=[debug_callback]
        )

        self.chat_history = []
        self.full_original_analysis_json = None
        self.current_formatting_plan_json = None
        self.full_modified_analysis_json = None
        self.current_doc_path = None
        self.current_output_doc_path = None
        self.PLACEHOLDER_ORIGINAL_ANALYSIS = "$FULL_ORIGINAL_ANALYSIS"
        self.PLACEHOLDER_MODIFIED_ANALYSIS = "$FULL_MODIFIED_ANALYSIS"
        self.PLACEHOLDER_FORMATTING_PLAN = "$CURRENT_FORMATTING_PLAN"

    def _initialize_llm(self):
        try:
            api_key = self.api_keys[self.current_key_index]
            self.llm = ChatGroq(
                groq_api_key=api_key,
                model_name=self.model_name,
                temperature=self.temperature
            )
            self._emit_progress({"type": "info", "message": f"LLM initialized with API key index {self.current_key_index}."})
        except Exception as e:
            self._emit_progress({"type": "error", "message": f"Fatal: Failed to initialize LLM with key index {self.current_key_index}: {e}"})
            raise ValueError(f"Failed to initialize LLM: {e}")

    def _rotate_api_key(self):
        self.current_key_index = (self.current_key_index + 1) % len(self.api_keys)
        self._emit_progress({"type": "warning", "message": f"Rotating API key to index {self.current_key_index}."})
        self._initialize_llm()
        # After rotating key, we need to re-create the agent with the new LLM instance
        react_agent = create_react_agent(llm=self.llm, tools=self.tools, prompt=self.prompt)
        
        # Recreate debug callback
        from langchain.callbacks.base import BaseCallbackHandler
        class DebugCallbackHandler(BaseCallbackHandler):
            def on_agent_action(self, action, **kwargs):
                print(f"DEBUG CALLBACK: Agent Action - Tool: {action.tool}, Input: {str(action.tool_input)[:200]}...")
            def on_agent_finish(self, finish, **kwargs):
                print(f"DEBUG CALLBACK: Agent Finish - Output: {str(finish.return_values)[:200]}...")
            def on_llm_end(self, response, **kwargs):
                print(f"DEBUG CALLBACK: LLM End - Response: {str(response.generations[0][0].text)[:200]}...")
        
        debug_callback = DebugCallbackHandler()
        
        # Update global reference for tools
        import smartdoc_agent.core.tools as tools_module
        tools_module._current_agent_instance = self
        
        self.agent = AgentExecutor(agent=react_agent, tools=self.tools, verbose=True, max_iterations=15, handle_parsing_errors=True, callbacks=[debug_callback])
        return self.current_key_index == 0 # Return True if we have tried all keys and wrapped around

    def _emit_progress(self, data: dict):
        """Helper to call the progress callback if it exists."""
        if self.progress_callback:
            try:
                self.progress_callback(data)
            except Exception as e:
                print(f"Error in progress_callback: {e}")

    def _substitute_placeholders(self, action_input: dict) -> dict:
        processed_input = {}
        for key, value in action_input.items():
            if value == self.PLACEHOLDER_ORIGINAL_ANALYSIS:
                if self.full_original_analysis_json:
                    processed_input[key] = self.full_original_analysis_json
                    self._emit_progress({"type": "debug", "message": f"Substituted {self.PLACEHOLDER_ORIGINAL_ANALYSIS} for key '{key}'"})
                else:
                    raise ValueError(f"Agent tried to use placeholder {self.PLACEHOLDER_ORIGINAL_ANALYSIS} but no original analysis is stored.")
            elif value == self.PLACEHOLDER_MODIFIED_ANALYSIS:
                if self.full_modified_analysis_json:
                    processed_input[key] = self.full_modified_analysis_json
                    self._emit_progress({"type": "debug", "message": f"Substituted {self.PLACEHOLDER_MODIFIED_ANALYSIS} for key '{key}'"})
                else:
                    raise ValueError(f"Agent tried to use placeholder {self.PLACEHOLDER_MODIFIED_ANALYSIS} but no modified analysis is stored.")
            elif value == self.PLACEHOLDER_FORMATTING_PLAN:
                if self.current_formatting_plan_json:
                    processed_input[key] = self.current_formatting_plan_json
                    self._emit_progress({"type": "debug", "message": f"Substituted {self.PLACEHOLDER_FORMATTING_PLAN} for key '{key}'"})
                else:
                    raise ValueError(f"Agent tried to use placeholder {self.PLACEHOLDER_FORMATTING_PLAN} but no plan is stored.")
            else:
                processed_input[key] = value
        return processed_input

    def _execute_tool(self, tool_name: str, tool_input_from_llm: Any) -> str:
        self._emit_progress({"type": "tool_start", "name": tool_name, "input_preview": str(tool_input_from_llm)[:100] + "..."})
        payload_value_for_tool: Any = None

        if isinstance(tool_input_from_llm, str):
                    parsed_llm_input_dict = None
                    try:
                        # ADD THIS LINE to fix unescaped backslashes from Windows paths
                        sanitized_input_str = tool_input_from_llm.replace('\\', '\\\\')
                        
                        # Clean up common JSON issues like trailing commas
                        import re
                        sanitized_input_str = re.sub(r',(\s*[}\]])', r'\1', sanitized_input_str)
                        
                        # Handle truncated JSON by ensuring it ends properly
                        if len(sanitized_input_str) > 10000:  # If JSON is very large
                            self._emit_progress({"type": "warning", "message": f"Large JSON input detected ({len(sanitized_input_str)} chars), attempting to truncate safely"})
                            # Try to find a safe truncation point
                            if '"user_goal"' in sanitized_input_str:
                                # Find the user_goal and try to create a minimal valid JSON
                                user_goal_match = re.search(r'"user_goal":\s*"([^"]*)"', sanitized_input_str)
                                if user_goal_match:
                                    user_goal = user_goal_match.group(1)
                                    sanitized_input_str = f'{{"user_goal": "{user_goal}"}}'
                                    self._emit_progress({"type": "info", "message": f"Simplified large JSON to minimal format with user_goal: {user_goal}"})

                        # UPDATE THIS LINE to use the new sanitized string
                        parsed_llm_input_dict = json.loads(sanitized_input_str)
                        
                        self._emit_progress({"type": "debug", "message": f"_execute_tool: Parsed LLM string input as JSON to dict."})
                        payload_value_for_tool = self._substitute_placeholders(parsed_llm_input_dict)
                    except json.JSONDecodeError:
                        payload_value_for_tool = tool_input_from_llm.strip('"').strip("'")
                        self._emit_progress({"type": "debug", "message": f"_execute_tool: LLM string input treated as simple string (JSON parse failed, after strip): {str(payload_value_for_tool)[:100]}..."})
                        
        elif isinstance(tool_input_from_llm, dict):
            payload_value_for_tool = self._substitute_placeholders(tool_input_from_llm)
            self._emit_progress({"type": "debug", "message": f"_execute_tool: LLM dict input. After subs: {str(payload_value_for_tool)[:100]}..."})
        else:
            error_msg = f"LLM provided invalid type for Action Input: {type(tool_input_from_llm)}. Value: {str(tool_input_from_llm)[:200]}"
            self._emit_progress({"type": "error", "message": f"_execute_tool: {error_msg}"})
            return json.dumps({"error": error_msg})

        args_for_langchain_invoke = {"tool_input": payload_value_for_tool}

        selected_tool = next((t for t in self.tools if t.name == tool_name), None)
        if not selected_tool:
            error_msg = f"Error: Tool '{tool_name}' not found. Available tools: {[t.name for t in self.tools]}"
            self._emit_progress({"type": "error", "message": error_msg})
            return error_msg

        observation = ""
        try:
            self._emit_progress({"type": "debug", "message": f"Invoking tool '{selected_tool.name}' with args for invoke: {str(args_for_langchain_invoke)[:200]}..."})
            observation = selected_tool.invoke(args_for_langchain_invoke)
            self._emit_progress({"type": "tool_end", "name": tool_name, "observation_preview": str(observation)[:100] + "..."})
        except Exception as e:
            error_msg_tool_exec = f"Error executing tool {tool_name}: {str(e)}"
            self._emit_progress({"type": "error", "message": error_msg_tool_exec, "details": traceback.format_exc()})
            traceback.print_exc() # Also print to server log for immediate visibility
            return error_msg_tool_exec # Return error message as observation

        try:
            # Attempt to parse observation to see if it's JSON (many tools return JSON strings)
            # This is for storing specific data structures, not for general logging.
            parsed_observation_for_storage = json.loads(observation) if isinstance(observation, str) else observation

            if tool_name == "analyze_document_structure" and isinstance(parsed_observation_for_storage, dict) and "document_path" in parsed_observation_for_storage:
                analyzed_path = parsed_observation_for_storage["document_path"]
                if analyzed_path == self.current_doc_path:
                    self.full_original_analysis_json = observation
                    self._emit_progress({"type": "data_store", "variable": "full_original_analysis_json", "length": len(observation) if observation else 0})
                elif analyzed_path == self.current_output_doc_path:
                    self.full_modified_analysis_json = observation
                    self._emit_progress({"type": "data_store", "variable": "full_modified_analysis_json", "length": len(observation) if observation else 0})
            elif tool_name == "create_formatting_plan":
                # Ensure what's stored is the JSON string if the tool returned a string
                if isinstance(observation, str):
                    self.current_formatting_plan_json = observation
                    self._emit_progress({"type": "data_store", "variable": "current_formatting_plan_json", "length": len(observation) if observation else 0})
                elif isinstance(parsed_observation_for_storage, (dict, list)): # If tool already parsed it (should not based on current tools.py)
                    self.current_formatting_plan_json = json.dumps(parsed_observation_for_storage) # Store as string
                    self._emit_progress({"type": "data_store", "variable": "current_formatting_plan_json", "length": len(self.current_formatting_plan_json)})
            elif tool_name == "apply_contextual_formatting":
                # Check if formatting was successful by parsing the observation
                try:
                    formatting_result = json.loads(observation) if isinstance(observation, str) else observation
                    if isinstance(formatting_result, dict):
                        if formatting_result.get("status") == "success" and formatting_result.get("output_file_exists"):
                            self._emit_progress({"type": "info", "message": "Formatting applied successfully. Output file verified to exist."})
                        elif formatting_result.get("status") == "error":
                            self._emit_progress({"type": "warning", "message": f"Formatting completed with issues: {formatting_result.get('message', 'Unknown error')}"})
                        else:
                            self._emit_progress({"type": "warning", "message": "Formatting status unclear from tool response."})
                    else:
                        self._emit_progress({"type": "info", "message": "Formatting applied. A subsequent call to 'analyze_document_structure' on the output path will be needed for validation."})
                except (json.JSONDecodeError, Exception) as e:
                    self._emit_progress({"type": "warning", "message": f"Could not parse formatting result: {e}. Assuming formatting completed."})
        except json.JSONDecodeError:
            self._emit_progress({"type":"debug", "message":f"Observation from {tool_name} was not valid JSON, not performing special storage: {observation[:100]}"})
        except Exception as e: # Catch other errors during storage logic
            self._emit_progress({"type":"error", "message":f"Error processing/storing observation from {tool_name}: {e}"})

        return str(observation)

    def run(self, user_query: str, document_path: str, output_document_path: str = None):
        self._emit_progress({"type": "lifecycle", "event": "agent_run_start", "user_query": user_query, "doc_path": document_path})
        self.current_doc_path = document_path
        if not output_document_path:
            self.current_output_doc_path = document_path.replace(".docx", "_formatted.docx")
        else:
            self.current_output_doc_path = output_document_path

        self.full_original_analysis_json = None
        self.current_formatting_plan_json = None
        self.full_modified_analysis_json = None

        agent_initial_input = (
            f"User query: '{user_query}'.\n"
            f"The document to work on is: '{self.current_doc_path}'.\n"
            f"The output path for formatted document is: '{self.current_output_doc_path}'.\n"
            "Follow the general workflow: analyze original, create plan, apply plan, analyze modified, validate. Use placeholders like $FULL_ORIGINAL_ANALYSIS where instructed."
        )

        if not self.chat_history:
            self.chat_history.append(HumanMessage(content=agent_initial_input))
        else:
            self.chat_history.append(HumanMessage(content=agent_initial_input))

        intermediate_steps = []
        max_iterations = 12
        attempts_with_current_key = 0

        # AgentExecutor handles the entire execution loop internally
        try:
            self._emit_progress({"type": "llm_call_start", "purpose": "agent_execution", "input_preview": agent_initial_input[:100]+"..."})
            
            # DEBUG: Print what we're sending to the agent
            print(f"\n=== AGENT EXECUTOR INPUT DEBUG ===")
            print(f"Input: {agent_initial_input}")
            print(f"Chat history count: {len(self.chat_history)}")
            print("=" * 50)
            
            agent_input = {
                "input": agent_initial_input,
                "chat_history": self.chat_history
            }
            
            llm_call_start_time = time.time()
            print(f"DEBUG: About to invoke AgentExecutor with API key index {self.current_key_index}")
            result = self.agent.invoke(agent_input)
            llm_call_duration = time.time() - llm_call_start_time
            print(f"DEBUG: AgentExecutor completed successfully")
            
            # DEBUG: Print what the agent returned
            print(f"\n=== AGENT EXECUTOR OUTPUT DEBUG ===")
            print(f"Output Type: {type(result)}")
            print(f"Output Content: {str(result)[:1000]}...")
            print("=" * 50)
            
            self._emit_progress({"type": "llm_call_end", "duration_seconds": round(llm_call_duration,2) , "output_type": str(type(result))})

            # AgentExecutor returns a dict with 'output' key when finished
            if isinstance(result, dict) and 'output' in result:
                final_answer = result['output']
                self._emit_progress({"type": "agent_finish", "final_answer": final_answer})
                self.chat_history.append(AIMessage(content=final_answer))
                return final_answer
            else:
                # This shouldn't happen with AgentExecutor
                error_msg = f"Unexpected result from AgentExecutor: {type(result)}, Content: {str(result)[:500]}"
                self._emit_progress({"type": "error", "message": error_msg})
                return f"Agent error: {error_msg}"

        except APIError as e:
            self._emit_progress({"type": "error", "message": f"API Error with key index {self.current_key_index}: {e}"})
            print(f"DEBUG: API Error details: {str(e)}")
            print(f"DEBUG: Error type: {type(e)}")
            print(f"DEBUG: Current API key (masked): {self.api_keys[self.current_key_index][:10]}...")
            
            all_keys_tried = self._rotate_api_key()
            if all_keys_tried:
                error_message = "All API keys have failed. Aborting."
                self._emit_progress({"type": "error", "message": error_message})
                self.chat_history.append(AIMessage(content=error_message))
                return error_message
            else:
                # Retry with new key - but avoid infinite recursion
                self._emit_progress({"type": "info", "message": f"Retrying with new API key (index {self.current_key_index})"})
                # Clear chat history to avoid accumulation on retry
                retry_chat_history = self.chat_history.copy()
                self.chat_history = []
                try:
                    return self.run(user_query, document_path, output_document_path)
                except Exception as retry_e:
                    # If retry fails, restore chat history and return error
                    self.chat_history = retry_chat_history
                    error_message = f"Retry failed: {str(retry_e)}"
                    self._emit_progress({"type": "error", "message": error_message})
                    return error_message

        except Exception as e:
            error_message = f"Unexpected error during agent execution: {str(e)}"
            self._emit_progress({"type": "error", "message": error_message, "details": traceback.format_exc()})
            self.chat_history.append(AIMessage(content=error_message))
            return error_message


if __name__ == '__main__':
    from docx import Document as PythonDocXDocument
    sample_doc_path = "sample_doc_for_agent.docx"
    sample_output_path = "sample_doc_for_agent_formatted.docx"

    def cli_progress_callback(data: dict):
        print(f"CLI_PROGRESS: {json.dumps(data)}")

    doc = PythonDocXDocument()
    doc.add_heading('Agent Test Document', level=0)
    doc.add_paragraph('This is paragraph one. It needs some formatting.')
    doc.add_heading('Section Alpha', level=1)
    doc.add_paragraph('This is paragraph two, under Section Alpha. It is a bit messy.')
    doc.save(sample_doc_path)
    print(f"Created '{sample_doc_path}' for agent testing.")

    try:
        formatter_agent = DocumentFormattingAgent(progress_callback=cli_progress_callback)

        user_request = "Please make this document look more professional and fix inconsistencies."
        print(f"\n--- Running agent with query: '{user_request}' on '{sample_doc_path}' ---")

        result = formatter_agent.run(
            user_query=user_request,
            document_path=sample_doc_path,
            output_document_path=sample_output_path
        )

        print("\n--- Agent Result ---")
        print(result)

    except ValueError as ve:
        print(f"Setup Error: {ve}")
        if "GROQ_API_KEY is not set" in str(ve):
             print("Please ensure your GROQ_API_KEY is set in .env at the project root.")
    except Exception as ex:
        print(f"An unexpected error occurred: {ex}")
        traceback.print_exc()
    finally:
        import os
        if os.path.exists(sample_doc_path):
            os.remove(sample_doc_path)
        if os.path.exists(sample_output_path):
            os.remove(sample_output_path)
        print(f"\nCleaned up '{sample_doc_path}' and '{sample_output_path}'.")
