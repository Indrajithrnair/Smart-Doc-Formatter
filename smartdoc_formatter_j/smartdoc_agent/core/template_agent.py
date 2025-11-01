"""
Template Formatting Agent
Extracts content from raw documents and fills business proposal templates.
"""

from langchain_groq import ChatGroq
from docx import Document
import json
import re
import os
from datetime import datetime
from smartdoc_agent.config import get_groq_api_keys
from typing import Dict, Optional


class TemplateFormattingAgent:
    """
    Agent specialized in converting unstructured documents to business proposal templates.
    Uses LLM to extract and structure content intelligently.
    """
    
    def __init__(self, 
                 template_name: str = "Business_Proposal_Template_With_Placeholders",
                 model_name: str = "llama-3.3-70b-versatile",
                 temperature: float = 0.3,
                 progress_callback=None):
        """
        Initialize the Template Formatting Agent.
        
        Args:
            template_name: Name of template file (without .docx extension)
            model_name: LLM model to use for content extraction
            temperature: LLM temperature (0.3 = slightly creative but focused)
            progress_callback: Callback function for progress updates
        """
        self.template_name = template_name
        self.progress_callback = progress_callback
        
        # Set up template path
        current_dir = os.path.dirname(os.path.abspath(__file__))
        self.template_path = os.path.join(
            current_dir, 
            "..", 
            "templates", 
            f"{template_name}.docx"
        )
        
        # Initialize LLM
        try:
            api_keys = get_groq_api_keys()
            self.llm = ChatGroq(
                api_key=api_keys[0],
                model_name=model_name,
                temperature=temperature
            )
            self._emit_progress("info", "Template Agent initialized with LLM")
        except Exception as e:
            self._emit_progress("error", f"Failed to initialize LLM: {e}")
            raise
        
        # Verify template exists
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found: {self.template_path}")
    
    def _emit_progress(self, event_type: str, message: str, data: dict = None):
        """Emit progress update if callback is provided"""
        if self.progress_callback:
            progress_data = {
                "type": event_type,
                "message": message
            }
            if data:
                progress_data.update(data)
            self.progress_callback(progress_data)
    
    def run(self, 
            document_path: str, 
            output_document_path: str,
            template_type: str = "business_proposal") -> str:
        """
        Main execution flow: Extract content and fill template.
        
        Args:
            document_path: Path to input document
            output_document_path: Path where formatted document will be saved
            template_type: Type of template (business_proposal or course_plan)
        
        Returns:
            Success message with output path
        """
        try:
            self._emit_progress("lifecycle", f"Starting {template_type} template formatting")
            
            # Step 1: Extract raw content from document
            self._emit_progress("step", "Reading input document...")
            raw_content = self._read_document(document_path)
            
            # Step 2: Extract structured content using LLM (template-specific)
            self._emit_progress("step", "Analyzing content with AI...")
            if template_type == "course_plan":
                extracted_content = self._extract_course_plan_content(raw_content)
            else:
                extracted_content = self._extract_business_proposal_content(raw_content)
            
            # Step 3: Load template
            template_name = "PMC2423-Course Plan_Template_With_Placeholders" if template_type == "course_plan" else self.template_name
            template_path = os.path.join(
                os.path.dirname(self.template_path),
                f"{template_name}.docx"
            )
            self._emit_progress("step", f"Loading {template_type} template...")
            template_doc = Document(template_path)
            
            # Step 4: Fill template with extracted content
            self._emit_progress("step", "Filling template with extracted content...")
            self._fill_template(template_doc, extracted_content, template_type)
            
            # Step 5: Save result
            self._emit_progress("step", "Saving formatted document...")
            template_doc.save(output_document_path)
            
            self._emit_progress("lifecycle", "Template formatting completed successfully")
            
            return f"Successfully created {template_type}: {output_document_path}"
            
        except Exception as e:
            error_msg = f"Template formatting failed: {str(e)}"
            self._emit_progress("error", error_msg)
            raise Exception(error_msg)
    
    def _read_document(self, document_path: str) -> str:
        """Read all text content from the input document"""
        try:
            doc = Document(document_path)
            
            # Extract all text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            raw_content = "\n\n".join(paragraphs)
            
            self._emit_progress("info", f"Read {len(paragraphs)} sections from document")
            return raw_content
            
        except Exception as e:
            raise Exception(f"Failed to read document: {e}")
    
    def _extract_business_proposal_content(self, raw_content: str) -> Dict[str, str]:
        """
        Use LLM to extract and structure content for business proposal template.
        Maps raw content to the 12 template placeholders.
        """
        prompt = f"""
You are an expert business analyst. Analyze the following document content and extract structured information for a business proposal.

Document Content:
{raw_content}

Extract the following sections and return them as a JSON object. For each field:
- Extract relevant content from the document
- If content is too brief, expand it professionally
- If a section is missing, create appropriate professional content based on context
- Keep language professional and business-appropriate

Required fields:
1. "company_name": The company or organization name (extract or infer)
2. "prepared_by": Who prepared this (extract or use "Professional Services Team")
3. "executive_summary": A compelling 2-3 sentence overview of the proposal
4. "problem_statement": Clear description of the problem/need being addressed
5. "proposed_solution": Detailed solution and approach (3-5 sentences)
6. "objectives_and_deliverables": Bulleted list of specific objectives and what will be delivered
7. "implementation_details": Technical or methodological details of how solution will be implemented
8. "timeline_and_milestones": Project timeline with key milestones (if available, otherwise suggest a reasonable timeline)
9. "budget_and_costs": Pricing information or investment details (if available, otherwise indicate "To be discussed")
10. "conclusion": Strong closing paragraph with call to action
11. "references": Case studies, testimonials, or relevant references (if available, otherwise use "References available upon request")

Return ONLY a valid JSON object with these exact keys. No markdown formatting, no code blocks, just the JSON.

Example format:
{{
    "company_name": "Acme Corporation",
    "prepared_by": "Strategic Solutions Team",
    "executive_summary": "This proposal outlines...",
    ...
}}
"""
        
        try:
            response = self.llm.invoke(prompt)
            content = response.content.strip()
            
            # Clean up response (remove markdown code blocks if present)
            content = re.sub(r'^```json\s*', '', content)
            content = re.sub(r'\s*```$', '', content)
            content = content.strip()
            
            # Parse JSON
            extracted_data = json.loads(content)
            
            # Add current date
            extracted_data['date'] = datetime.now().strftime("%B %d, %Y")
            
            # Log what was extracted
            self._emit_progress("info", f"Extracted {len(extracted_data)} sections from content")
            
            return extracted_data
            
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse LLM response as JSON: {e}\nResponse: {content[:200]}")
        except Exception as e:
            raise Exception(f"Content extraction failed: {e}")
    
    def _extract_course_plan_content(self, raw_content: str) -> Dict[str, str]:
        """
        Use LLM to extract comprehensive course plan data from syllabus.
        Handles 100+ placeholders with intelligent mapping.
        """
        self._emit_progress("info", "Extracting course structure from syllabus...")
        
        # Step 1: Extract basic course info and outcomes
        basic_prompt = f"""
You are an expert curriculum designer. Analyze this course syllabus and extract structured data.

Syllabus Content:
{raw_content}

Extract and return as JSON:
{{
    "course_name": "Full course name",
    "course_code": "Course code (e.g., PMC2418)",
    "academic_level": "Academic level (e.g., UG, PG, Diploma)",
    "course_type": "Core/Elective/etc",
    "faculty_name": "Faculty name if mentioned, else 'Faculty Name'",
    "course_description": "Detailed 2-3 sentence description",
    "prerequisites": "Prerequisites or 'None'",
    "course_credit": "Total course credits (e.g., 3)",
    "theory_credit": "Theory credits/hours (e.g., 3)",
    "practical_credit": "Practical/Lab credits/hours (e.g., 0)",
    "total_credit": "Total credits (theory + practical)",
    "credits": "Number of credits",
    "total_hours": "Total instructional hours",
    "co1": "First course outcome verbatim",
    "co2": "Second course outcome verbatim",
    "co3": "Third course outcome verbatim",
    "co4": "Fourth course outcome verbatim",
    "co5": "Fifth course outcome verbatim",
    "modules": [
        {{
            "number": 1,
            "title": "Module title",
            "hours": 14,
            "topics": ["topic1", "topic2", "..."]
        }},
        ...
    ]
}}

Return ONLY valid JSON, no markdown.
"""
        
        try:
            response = self.llm.invoke(basic_prompt)
            basic_content = self._clean_json_response(response.content)
            
            # Debug: Print what we got from LLM
            print(f"[DEBUG] LLM Response length: {len(basic_content)}")
            print(f"[DEBUG] LLM Response (first 500 chars): {basic_content[:500]}")
            
            if not basic_content or basic_content.strip() == "":
                raise ValueError("LLM returned empty response")
            
            basic_data = json.loads(basic_content)
            
            self._emit_progress("info", f"Extracted basic info and {len(basic_data.get('modules', []))} modules")
            
            # Step 2: Generate micro-outcomes for each CO
            self._emit_progress("info", "Generating micro-outcomes for each CO...")
            micro_outcomes = self._generate_micro_outcomes(basic_data)
            
            # Step 3: Map modules to micro-outcomes and generate detailed blueprint
            self._emit_progress("info", "Creating course blueprint with module mappings...")
            blueprint = self._create_course_blueprint(basic_data, micro_outcomes)
            
            # Step 4: Generate assessment schemes
            self._emit_progress("info", "Generating assessment schemes...")
            assessments = self._generate_assessment_schemes(basic_data)
            
            # Step 5: Compile all data
            complete_data = {
                **basic_data,
                **micro_outcomes,
                **blueprint,
                **assessments,
                'date': datetime.now().strftime("%B %d, %Y")
            }
            
            self._emit_progress("info", f"Course plan extraction complete: {len(complete_data)} fields populated")
            return complete_data
            
        except json.JSONDecodeError as e:
            error_msg = f"JSON parsing failed. LLM response was: '{basic_content[:200]}...'"
            print(f"[ERROR] {error_msg}")
            self._emit_progress("error", error_msg)
            raise Exception(f"Course plan extraction failed: Invalid JSON from LLM - {e}")
        except Exception as e:
            print(f"[ERROR] Course plan extraction error: {e}")
            self._emit_progress("error", f"Extraction failed: {str(e)}")
            raise Exception(f"Course plan extraction failed: {e}")
    
    def _clean_json_response(self, content: str) -> str:
        """Clean LLM response to extract pure JSON"""
        if not content:
            return ""
        
        content = content.strip()
        
        # Remove markdown code blocks
        content = re.sub(r'^```json\s*', '', content)
        content = re.sub(r'^```\s*', '', content)
        content = re.sub(r'\s*```$', '', content)
        
        # Find JSON object if embedded in text
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            content = json_match.group(0)
        
        return content.strip()
    
    def _generate_micro_outcomes(self, basic_data: Dict) -> Dict[str, str]:
        """Generate 3 micro-outcomes for each of the 5 COs"""
        prompt = f"""
Given these 5 course outcomes, generate 3 specific, measurable micro-outcomes for each.

CO1: {basic_data.get('co1', '')}
CO2: {basic_data.get('co2', '')}
CO3: {basic_data.get('co3', '')}
CO4: {basic_data.get('co4', '')}
CO5: {basic_data.get('co5', '')}

For each CO, create 3 micro-outcomes that are:
- Specific and measurable
- Progressive in complexity
- Aligned with the main CO

Return as JSON:
{{
    "co1_micro_1": "CO1.1: Specific micro-outcome...",
    "co1_micro_2": "CO1.2: Next micro-outcome...",
    "co1_micro_3": "CO1.3: Advanced micro-outcome...",
    "co2_micro_1": "CO2.1: ...",
    ... (repeat for all 5 COs)
}}

Return ONLY valid JSON.
"""
        
        response = self.llm.invoke(prompt)
        content = self._clean_json_response(response.content)
        return json.loads(content)
    
    def _create_course_blueprint(self, basic_data: Dict, micro_outcomes: Dict) -> Dict[str, str]:
        """Map modules to micro-outcomes and generate complete blueprint"""
        modules_text = "\n".join([
            f"Module {m['number']}: {m['title']} ({m['hours']} hours) - Topics: {', '.join(m.get('topics', []))}"
            for m in basic_data.get('modules', [])
        ])
        
        prompt = f"""
Create a detailed course blueprint mapping modules to micro-outcomes.

Course Outcomes and Micro-outcomes:
{json.dumps(micro_outcomes, indent=2)}

Modules:
{modules_text}

For each of the 15 micro-outcomes (3 per CO), assign:
1. module: Which module covers this (e.g., "Module 1")
2. total_hours: Hours for this micro-outcome (distribute module hours)
3. division: How hours are divided (e.g., "2L + 1T")
4. topic: Specific topic from module
5. tla: Teaching-Learning Activity (lecture/lab/discussion/project/etc)
6. tools: Tools/resources (IDE/LMS/Whiteboard/etc)
7. assessment: Assessment method (Quiz/Assignment/Presentation/etc)
8. bloom: Bloom's level (Remember/Understand/Apply/Analyze/Evaluate/Create)

Return as JSON with keys like:
{{
    "co1_module_1": "Module 1",
    "co1_total_hours_1": "5",
    "co1_division_1": "3L + 2P",
    "co1_topic_1": "Development Environment Setup",
    "co1_tla_1": "Hands-on Lab",
    "co1_tools_1": "IntelliJ IDEA, JDK",
    "co1_assessment_1": "Lab Exercise",
    "co1_bloom_1": "Apply",
    ... (repeat for all 15 micro-outcomes: co1_1, co1_2, co1_3, co2_1, ...)
}}

Return ONLY valid JSON.
"""
        
        response = self.llm.invoke(prompt)
        content = self._clean_json_response(response.content)
        return json.loads(content)
    
    def _generate_assessment_schemes(self, basic_data: Dict) -> Dict[str, str]:
        """Generate assessment scheme details"""
        assessment_data = {
            # ISA (Internal Summative Assessment)
            'isa_marks': '20',
            'isa_description': 'Mid-term examination covering modules 1-3',
            'isa_co1': '4', 'isa_co2': '4', 'isa_co3': '4', 'isa_co4': '4', 'isa_co5': '4',
            
            # LB (Lab/Practical)
            'lb_marks': '10',
            'lb_description': 'Practical assignments and lab exercises',
            'lb_co1': '2', 'lb_co2': '2', 'lb_co3': '2', 'lb_co4': '2', 'lb_co5': '2',
            
            # LD (Learning Diary/Project)
            'ld_marks': '5',
            'ld_description': 'Project work and continuous documentation',
            'ld_co1': '1', 'ld_co2': '1', 'ld_co3': '1', 'ld_co4': '1', 'ld_co5': '1',
            
            # CP (Class Participation)
            'cp_marks': '5',
            'cp_description': 'Active participation, attendance, and engagement',
            
            # SEA1 (Semester End Assessment 1)
            'sea1_marks': '40',
            'sea1_description': 'Comprehensive final examination - Part A',
            'sea1_co1': '8', 'sea1_co2': '8', 'sea1_co3': '8', 'sea1_co4': '8', 'sea1_co5': '8',
            
            # SEA2 (Semester End Assessment 2)
            'sea2_marks': '20',
            'sea2_description': 'Comprehensive final examination - Part B',
            'sea2_co1': '4', 'sea2_co2': '4', 'sea2_co3': '4', 'sea2_co4': '4', 'sea2_co5': '4',
            
            # Other fields
            'delivery_method': 'Blended learning with lectures, labs, and flipped classroom',
            'lms_details': 'Moodle LMS for content delivery, assignments, and assessments',
            'indirect_assessment_method': 'Course feedback surveys and exit questionnaires'
        }
        
        return assessment_data
    
    def _build_course_plan_mapping(self, content: Dict) -> Dict[str, str]:
        """Build comprehensive placeholder mapping for course plan template (100+ placeholders)"""
        mapping = {}
        
        # Basic course info
        mapping['{{COURSE_NAME}}'] = content.get('course_name', '[Course Name]')
        mapping['{{COURSE_CODE}}'] = content.get('course_code', '[Code]')
        mapping['{{ACADEMIC_LEVEL}}'] = content.get('academic_level', 'UG')
        mapping['{{COURSE_TYPE}}'] = content.get('course_type', 'Core')
        mapping['{{FACULTY_NAME}}'] = content.get('faculty_name', '[Faculty Name]')
        mapping['{{COURSE_DESCRIPTION}}'] = content.get('course_description', '')
        mapping['{{PREREQUISITES}}'] = content.get('prerequisites', 'None')
        
        # Credit information
        mapping['{{COURSE_CREDIT}}'] = content.get('course_credit', '3')
        mapping['{{THEORY_CREDIT}}'] = content.get('theory_credit', '3')
        mapping['{{PRACTICAL_CREDIT}}'] = content.get('practical_credit', '0')
        mapping['{{TOTAL_CREDIT}}'] = content.get('total_credit', '3')
        
        # Course Outcomes (5 main COs)
        for i in range(1, 6):
            mapping[f'{{{{CO{i}}}}}'] = content.get(f'co{i}', f'[Course Outcome {i}]')
        
        # Micro-outcomes (3 per CO = 15 total)
        for co_num in range(1, 6):
            for micro_num in range(1, 4):
                key = f'co{co_num}_micro_{micro_num}'
                mapping[f'{{{{CO{co_num}_MICRO_{micro_num}}}}}'] = content.get(key, f'CO{co_num}.{micro_num}: [Micro-outcome]')
        
        # Blueprint details for each micro-outcome (15 micro-outcomes × 8 fields each = 120 fields)
        for co_num in range(1, 6):
            for micro_num in range(1, 4):
                prefix = f'co{co_num}'
                # Module
                mapping[f'{{{{CO{co_num}_MODULE_{micro_num}}}}}'] = content.get(f'{prefix}_module_{micro_num}', 'Module 1')
                # Total Hours
                mapping[f'{{{{CO{co_num}_TOTAL_HOURS_{micro_num}}}}}'] = content.get(f'{prefix}_total_hours_{micro_num}', '5')
                # Division of Hours
                mapping[f'{{{{CO{co_num}_DIVISION_{micro_num}}}}}'] = content.get(f'{prefix}_division_{micro_num}', '3L + 2P')
                # Topic
                mapping[f'{{{{CO{co_num}_TOPIC_{micro_num}}}}}'] = content.get(f'{prefix}_topic_{micro_num}', '[Topic]')
                # Teaching-Learning Activity
                mapping[f'{{{{CO{co_num}_TLA_{micro_num}}}}}'] = content.get(f'{prefix}_tla_{micro_num}', 'Lecture')
                # Tools
                mapping[f'{{{{CO{co_num}_TOOLS_{micro_num}}}}}'] = content.get(f'{prefix}_tools_{micro_num}', 'LMS, Whiteboard')
                # Assessment
                mapping[f'{{{{CO{co_num}_ASSESSMENT_{micro_num}}}}}'] = content.get(f'{prefix}_assessment_{micro_num}', 'Quiz')
                # Bloom's Taxonomy
                mapping[f'{{{{CO{co_num}_BLOOM_{micro_num}}}}}'] = content.get(f'{prefix}_bloom_{micro_num}', 'Understand')
        
        # Assessment Schemes
        # ISA (Internal Summative Assessment)
        mapping['{{ISA_MARKS}}'] = content.get('isa_marks', '20')
        mapping['{{ISA_DESCRIPTION}}'] = content.get('isa_description', 'Mid-term examination')
        for i in range(1, 6):
            mapping[f'{{{{ISA_CO{i}}}}}'] = content.get(f'isa_co{i}', '4')
        
        # LB (Lab/Practical)
        mapping['{{LB_MARKS}}'] = content.get('lb_marks', '10')
        mapping['{{LB_DESCRIPTION}}'] = content.get('lb_description', 'Lab exercises')
        for i in range(1, 6):
            mapping[f'{{{{LB_CO{i}}}}}'] = content.get(f'lb_co{i}', '2')
        
        # LD (Learning Diary)
        mapping['{{LD_MARKS}}'] = content.get('ld_marks', '5')
        mapping['{{LD_DESCRIPTION}}'] = content.get('ld_description', 'Project work')
        for i in range(1, 6):
            mapping[f'{{{{LD_CO{i}}}}}'] = content.get(f'ld_co{i}', '1')
        
        # CP (Class Participation)
        mapping['{{CP_MARKS}}'] = content.get('cp_marks', '5')
        mapping['{{CP_DESCRIPTION}}'] = content.get('cp_description', 'Participation and attendance')
        
        # SEA1 (Semester End Assessment 1)
        mapping['{{SEA1_MARKS}}'] = content.get('sea1_marks', '40')
        mapping['{{SEA1_DESCRIPTION}}'] = content.get('sea1_description', 'Final exam - Part A')
        for i in range(1, 6):
            mapping[f'{{{{SEA1_CO{i}}}}}'] = content.get(f'sea1_co{i}', '8')
        
        # SEA2 (Semester End Assessment 2)
        mapping['{{SEA2_MARKS}}'] = content.get('sea2_marks', '20')
        mapping['{{SEA2_DESCRIPTION}}'] = content.get('sea2_description', 'Final exam - Part B')
        for i in range(1, 6):
            mapping[f'{{{{SEA2_CO{i}}}}}'] = content.get(f'sea2_co{i}', '4')
        
        # Delivery and Assessment Methods
        mapping['{{DELIVERY_METHOD}}'] = content.get('delivery_method', 'Blended learning')
        mapping['{{LMS_DETAILS}}'] = content.get('lms_details', 'Moodle LMS')
        mapping['{{INDIRECT_ASSESSMENT_METHOD}}'] = content.get('indirect_assessment_method', 'Course feedback surveys')
        
        return mapping
    
    def _fill_template(self, template_doc: Document, content: Dict[str, str], template_type: str = "business_proposal"):
        """
        Fill template placeholders with extracted content.
        Handles paragraphs, tables, headers, and footers.
        """
        # Build placeholder mapping based on template type
        if template_type == "course_plan":
            placeholder_mapping = self._build_course_plan_mapping(content)
        else:
            placeholder_mapping = {
                '{{COMPANY_NAME}}': content.get('company_name', '[Company Name]'),
                '{{PREPARED_BY}}': content.get('prepared_by', '[Prepared By]'),
                '{{DATE}}': content.get('date', datetime.now().strftime("%B %d, %Y")),
                '{{EXECUTIVE_SUMMARY}}': content.get('executive_summary', ''),
                '{{PROBLEM_STATEMENT}}': content.get('problem_statement', ''),
                '{{PROPOSED_SOLUTION}}': content.get('proposed_solution', ''),
                '{{OBJECTIVES_AND_DELIVERABLES}}': content.get('objectives_and_deliverables', ''),
                '{{IMPLEMENTATION_DETAILS}}': content.get('implementation_details', ''),
                '{{TIMELINE_AND_MILESTONES}}': content.get('timeline_and_milestones', ''),
                '{{BUDGET_AND_COSTS}}': content.get('budget_and_costs', ''),
                '{{CONCLUSION}}': content.get('conclusion', ''),
                '{{REFERENCES}}': content.get('references', 'References available upon request')
            }
        
        replacements_made = 0
        
        # Replace in paragraphs
        for paragraph in template_doc.paragraphs:
            for placeholder, replacement in placeholder_mapping.items():
                if placeholder in paragraph.text:
                    # Replace while preserving formatting
                    self._replace_in_paragraph(paragraph, placeholder, replacement)
                    replacements_made += 1
        
        # Replace in tables
        for table in template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in placeholder_mapping.items():
                            if placeholder in paragraph.text:
                                self._replace_in_paragraph(paragraph, placeholder, replacement)
                                replacements_made += 1
        
        # Replace in headers
        for section in template_doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    for placeholder, replacement in placeholder_mapping.items():
                        if placeholder in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder, replacement)
                            replacements_made += 1
            
            # Replace in footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for placeholder, replacement in placeholder_mapping.items():
                        if placeholder in paragraph.text:
                            self._replace_in_paragraph(paragraph, placeholder, replacement)
                            replacements_made += 1
        
        self._emit_progress("info", f"Filled {replacements_made} placeholders in template")
    
    def _replace_in_paragraph(self, paragraph, placeholder: str, replacement: str):
        """
        Replace placeholder in paragraph while preserving formatting.
        Works with runs to maintain font, size, color, etc.
        """
        if placeholder not in paragraph.text:
            return
        
        # Simple approach: replace in full text
        # Note: This loses some formatting. For production, use run-level replacement
        full_text = paragraph.text
        new_text = full_text.replace(placeholder, str(replacement))
        
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''
        
        # Add new text (uses paragraph's base formatting)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def test_template_agent():
    """Test function for development"""
    print("Testing Template Formatting Agent...")
    
    # Create a test document
    test_input = "test_input.docx"
    test_output = "test_output_proposal.docx"
    
    # Create test input document
    doc = Document()
    doc.add_paragraph("""
    Acme Corporation needs a new customer management system. 
    We currently use spreadsheets which are inefficient.
    
    We propose building a custom CRM solution using modern cloud technology.
    The system will include contact management, sales pipeline tracking, and reporting.
    
    Timeline: 3 months
    Budget: $50,000
    
    Our team has 10 years of CRM development experience.
    """)
    doc.save(test_input)
    
    # Run agent
    agent = TemplateFormattingAgent()
    result = agent.run(test_input, test_output)
    
    print(f"\n✅ {result}")
    print(f"Output saved to: {test_output}")


if __name__ == "__main__":
    test_template_agent()
