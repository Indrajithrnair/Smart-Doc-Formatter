from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# --- Action Handlers ---

def apply_find_and_replace_font_action(doc: Document, elements_details: list, action: dict):
    """
    Finds all occurrences of a phrase and applies specific font formatting to them
    without replacing the text itself. This is a formatting-only find-and-replace.
    Action: {"action": "find_and_replace_font", "search_phrase": "CONFIDENTIAL", 
             "font_name": "Arial", "size": 12, "bold": true, "uppercase": true}
    
    NOTE: This implementation will reset other formatting (like italics or color) within
    any paragraph that contains the search_phrase.
    """
    print(f"Applying find_and_replace_font action: {action}")
    search_phrase = action.get("search_phrase")
    font_name = action.get("font_name")
    size_pt = action.get("size")
    bold = action.get("bold")
    is_uppercase = action.get("uppercase", False)

    if not search_phrase:
        print("Warning: 'search_phrase' not provided for find_and_replace_font. Skipping.")
        return

    pattern = re.compile(re.escape(search_phrase), re.IGNORECASE)
    
    modified_paras_count = 0
    for para in doc.paragraphs:
        if pattern.search(para.text):
            modified_paras_count += 1
            original_text = para.text
            para.text = ''
            last_end = 0
            for match in pattern.finditer(original_text):
                start, end = match.span()
                if start > last_end:
                    para.add_run(original_text[last_end:start])
                
                matched_text = original_text[start:end]
                if is_uppercase:
                    matched_text = matched_text.upper()
                
                formatted_run = para.add_run(matched_text)
                if font_name:
                    formatted_run.font.name = font_name
                if size_pt:
                    formatted_run.font.size = Pt(size_pt)
                if bold is not None:
                    formatted_run.bold = bold
                
                # Add color support to find_and_replace_font
                color_value = action.get("color")
                if color_value:
                    if color_value == "black":
                        formatted_run.font.color.rgb = RGBColor(0, 0, 0)
                    elif color_value == "white":
                        formatted_run.font.color.rgb = RGBColor(255, 255, 255)
                    elif color_value == "red":
                        formatted_run.font.color.rgb = RGBColor(255, 0, 0)
                    elif color_value == "blue":
                        formatted_run.font.color.rgb = RGBColor(0, 0, 255)
                    elif color_value == "green":
                        formatted_run.font.color.rgb = RGBColor(0, 128, 0)
                    elif isinstance(color_value, str) and color_value.startswith("#"):
                        # Handle hex colors like #000000
                        hex_color = color_value.lstrip("#")
                        if len(hex_color) == 6:
                            r = int(hex_color[0:2], 16)
                            g = int(hex_color[2:4], 16)
                            b = int(hex_color[4:6], 16)
                            formatted_run.font.color.rgb = RGBColor(r, g, b)
                
                last_end = end
            
            if last_end < len(original_text):
                para.add_run(original_text[last_end:])

    print(f"Applied find_and_replace_font formatting to {modified_paras_count} paragraphs.")

def apply_find_and_replace_action(doc: Document, elements_details: list, action: dict):
    print(f"Applying find_and_replace action: {action}")
    find_text = action.get("find")
    replace_with = action.get("replace_with")
    if not find_text or replace_with is None:
        print("Warning: 'find' or 'replace_with' not provided for find_and_replace. Skipping.")
        return
    replaced_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if find_text.lower() in run.text.lower():
                pattern = re.compile(re.escape(find_text), re.IGNORECASE)
                new_text, n = pattern.subn(replace_with, run.text)
                if n > 0:
                    run.text = new_text
                    replaced_count += n
    print(f"Applied find_and_replace to {replaced_count} occurrences of '{find_text}'.")

def apply_set_paragraph_spacing_action(doc: Document, elements_details: list, action: dict):
    print(f"Applying paragraph spacing action: {action}")
    scope = action.get("scope")
    spacing_before = action.get("spacing_before")
    spacing_after = action.get("spacing_after")
    line_spacing = action.get("line_spacing")

    target_paras = []
    if scope == "all_paragraphs" or scope == "all_body_paragraphs":
        # "all_paragraphs" is interpreted as body paragraphs only (excluding headings)
        # Use elements_details to filter out headings
        for el_detail in elements_details:
            if el_detail.get("type") == "paragraph" and el_detail["paragraph_index"] < len(doc.paragraphs):
                target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
    elif scope == "all_elements":
        # "all_elements" includes both paragraphs and headings
        target_paras = doc.paragraphs
    else:
        print(f"Warning: Unknown or unsupported scope '{scope}' for set_paragraph_spacing.")
        return

    for para in target_paras:
        set_paragraph_spacing_properties(para, spacing_before, spacing_after, line_spacing)
    print(f"Applied spacing to {len(target_paras)} paragraphs for scope '{scope}'.")

def apply_set_alignment_action(doc: Document, elements_details: list, action: dict):
    print(f"Applying alignment action: {action}")
    scope = action.get("scope")
    alignment = action.get("alignment")

    target_paras = []
    if scope == "all_paragraphs" or scope == "all_body_paragraphs":
        # "all_paragraphs" is interpreted as body paragraphs only (excluding headings)
        # Use elements_details to filter out headings
        for el_detail in elements_details:
            if el_detail.get("type") == "paragraph" and el_detail["paragraph_index"] < len(doc.paragraphs):
                target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
    elif scope == "all_elements":
        # "all_elements" includes both paragraphs and headings
        target_paras = doc.paragraphs
    elif scope and scope.startswith("headings_level_"):
        try:
            level = int(scope.split("_")[-1])
            for el_detail in elements_details:
                if el_detail.get("type") == "heading" and el_detail.get("level") == level and el_detail["paragraph_index"] < len(doc.paragraphs):
                    target_paras.append(doc.paragraphs[el_detail["paragraph_index"]])
        except ValueError:
            print(f"Warning: Invalid heading level in scope '{scope}'.")
    else:
        print(f"Warning: Unknown or unsupported scope '{scope}' for set_alignment.")
        return

    for para in target_paras:
        set_paragraph_alignment_properties(para, alignment)
    print(f"Applied alignment to {len(target_paras)} paragraphs for scope '{scope}'.")

def apply_fix_font_inconsistencies_action(doc: Document, elements_details: list, action: dict):
    print(f"Applying font inconsistency fix: {action}")
    target_font_name = action.get("target_font_name")
    target_font_size_pt = action.get("target_font_size")

    if not target_font_name and not target_font_size_pt:
        print("Warning: No target font name or size provided for fix_font_inconsistencies. Skipping.")
        return

    changed_elements_count = 0
    for para_idx, para_detail in enumerate(elements_details):
        if para_idx < len(doc.paragraphs):
            paragraph = doc.paragraphs[para_idx]
            for run in paragraph.runs:
                applied_change_to_run = False
                if target_font_name and run.font.name != target_font_name:
                    run.font.name = target_font_name
                    applied_change_to_run = True
                if target_font_size_pt and run.font.size != Pt(target_font_size_pt):
                    run.font.size = Pt(target_font_size_pt)
                    applied_change_to_run = True
                if applied_change_to_run:
                    changed_elements_count += 1
    print(f"Applied font inconsistency fix to {changed_elements_count} runs.")

def apply_set_heading_font_action(doc, elements_details, action):
    print(f"Applying action: set_heading_font with params {action}")
    level = action.get("level")
    exclude_level = action.get("exclude_level")
    doc_paragraphs = doc.paragraphs
    
    for element in elements_details:
        if element.get("type") == "heading":
            element_level = element.get("level")
            
            # Check if this heading should be excluded
            if exclude_level is not None and element_level == exclude_level:
                continue
            
            # Apply formatting if level matches or if level is "all"
            if (level == "all" or level == element_level) and element.get("paragraph_index") < len(doc_paragraphs):
                paragraph = doc_paragraphs[element.get("paragraph_index")]
                for run in paragraph.runs:
                    if "font_name" in action: run.font.name = action["font_name"]
                    if "size" in action: run.font.size = Pt(action["size"])
                    if "bold" in action: run.bold = action["bold"]
                    if "italic" in action: run.italic = action["italic"]
                    if "color" in action: 
                        color_value = action["color"]
                        if color_value == "black":
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        elif color_value == "white":
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        elif color_value == "red":
                            run.font.color.rgb = RGBColor(255, 0, 0)
                        elif color_value == "blue":
                            run.font.color.rgb = RGBColor(0, 0, 255)
                        elif color_value == "green":
                            run.font.color.rgb = RGBColor(0, 128, 0)
                        elif isinstance(color_value, str) and color_value.startswith("#"):
                            # Handle hex colors like #000000
                            hex_color = color_value.lstrip("#")
                            if len(hex_color) == 6:
                                r = int(hex_color[0:2], 16)
                                g = int(hex_color[2:4], 16)
                                b = int(hex_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)

def apply_set_paragraph_font_action(doc, elements_details, action):
    print(f"Applying action: set_paragraph_font with params {action}")
    doc_paragraphs = doc.paragraphs
    
    for element in elements_details:
        if element.get("type") == "paragraph" and element.get("paragraph_index") < len(doc_paragraphs):
            paragraph = doc_paragraphs[element.get("paragraph_index")]
            for run in paragraph.runs:
                if "font_name" in action: run.font.name = action["font_name"]
                if "size" in action: run.font.size = Pt(action["size"])
                if "bold" in action: run.bold = action["bold"]
                if "italic" in action: run.italic = action["italic"]
                if "color" in action: 
                    color_value = action["color"]
                    if color_value == "black":
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    elif color_value == "white":
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif color_value == "red":
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    elif color_value == "blue":
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    elif color_value == "green":
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif isinstance(color_value, str) and color_value.startswith("#"):
                        # Handle hex colors like #000000
                        hex_color = color_value.lstrip("#")
                        if len(hex_color) == 6:
                            r = int(hex_color[0:2], 16)
                            g = int(hex_color[2:4], 16)
                            b = int(hex_color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)

def apply_set_document_default_font_action(doc, elements_details, action):
    print(f"Applying action: set_document_default_font with params {action}")
    style = doc.styles['Normal']
    font = style.font
    
    if "font_name" in action: font.name = action["font_name"]
    if "size" in action: font.size = Pt(action["size"])
    if "bold" in action: font.bold = action["bold"]
    if "italic" in action: font.italic = action["italic"]
    if "color" in action: 
        color_value = action["color"]
        if color_value == "black":
            font.color.rgb = RGBColor(0, 0, 0)
        elif color_value == "white":
            font.color.rgb = RGBColor(255, 255, 255)
        elif color_value == "red":
            font.color.rgb = RGBColor(255, 0, 0)
        elif color_value == "blue":
            font.color.rgb = RGBColor(0, 0, 255)
        elif color_value == "green":
            font.color.rgb = RGBColor(0, 128, 0)
        elif isinstance(color_value, str) and color_value.startswith("#"):
            # Handle hex colors like #000000
            hex_color = color_value.lstrip("#")
            if len(hex_color) == 6:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)

# --- Core Utilities ---

def load_document(file_path: str) -> Document:
    try:
        return Document(file_path)
    except Exception as e:
        print(f"Error loading document {file_path}: {e}")
        raise

def save_document(document: Document, file_path: str) -> None:
    try:
        document.save(file_path)
        print(f"Document saved to {file_path}")
    except Exception as e:
        print(f"Error saving document to {file_path}: {e}")
        raise

def get_document_analysis(document: Document) -> dict:
    analysis = {"elements": [get_paragraph_details(para, i) for i, para in enumerate(document.paragraphs)]}
    return analysis

def get_paragraph_details(para, para_index: int) -> dict:
    alignment_name = None
    if para.alignment is not None:
        try:
            # The .name attribute exists on the WD_ALIGN_PARAGRAPH enum members
            alignment_name = para.alignment.name
        except AttributeError:
            # Fallback for other EnumValue types that might not have .name
            alignment_name = str(para.alignment)

    para_info = {
        "paragraph_index": para_index,
        "text": para.text,
        "style_name": para.style.name if para.style else "Default Paragraph Font",
        "alignment": alignment_name,
        "runs": [get_run_details(run) for run in para.runs]
    }

    if para.style and para.style.name.startswith('Heading'):
        try:
            para_info.update({"type": "heading", "level": int(para.style.name.split(' ')[-1])})
        except ValueError:
            para_info.update({"type": "paragraph", "level": 0})
    elif para.style and para.style.name == "Title":
        para_info.update({"type": "heading", "level": 0})
    else:
        para_info["type"] = "paragraph"

    return para_info

def get_run_details(run) -> dict:
    return {
        "text": run.text,
        "font_name": run.font.name,
        "font_size": run.font.size.pt if run.font.size else None,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
    }

def set_paragraph_font_properties(paragraph, font_name: str = None, size_pt: float = None, bold: bool = None, italic: bool = None, underline: bool = None, color: str = None):
    for run in paragraph.runs:
        if font_name: run.font.name = font_name
        if size_pt: run.font.size = Pt(size_pt)
        if bold is not None: run.bold = bold
        if italic is not None: run.italic = italic
        if underline is not None: run.underline = underline
        if color:
            if color == "black":
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif color == "white":
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif color == "red":
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif color == "blue":
                run.font.color.rgb = RGBColor(0, 0, 255)
            elif color == "green":
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif isinstance(color, str) and color.startswith("#"):
                # Handle hex colors like #000000
                hex_color = color.lstrip("#")
                if len(hex_color) == 6:
                    r = int(hex_color[0:2], 16)
                    g = int(hex_color[2:4], 16)
                    b = int(hex_color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)

def set_paragraph_spacing_properties(paragraph, spacing_before_pt: float = None, spacing_after_pt: float = None, line_spacing_rule: float = None):
    if spacing_before_pt is not None: paragraph.paragraph_format.space_before = Pt(spacing_before_pt)
    if spacing_after_pt is not None: paragraph.paragraph_format.space_after = Pt(spacing_after_pt)
    if line_spacing_rule is not None: paragraph.paragraph_format.line_spacing = line_spacing_rule

def set_paragraph_alignment_properties(paragraph, alignment: str = None):
    if alignment:
        try:
            align_enum = getattr(WD_ALIGN_PARAGRAPH, alignment.upper(), None)
            if align_enum is not None:
                paragraph.alignment = align_enum
            else:
                print(f"Warning: Invalid alignment value '{alignment}'. Skipping.")
        except Exception as e:
            print(f"Warning: Exception setting alignment '{alignment}': {e}")

# --- NEW: Content Transformation Actions ---

def apply_markdown_to_headings_action(doc, elements_details, action):
    """
    Converts markdown-style headings to Word headings
    Action: {"action": "convert_markdown_to_headings", "single_hash_level": 0, "double_hash_level": 1}
    - Lines starting with # become Title (level 0) or Heading 1 (level 1)
    - Lines starting with ## become Heading 1 or Heading 2
    """
    print(f"Applying markdown to headings conversion: {action}")
    
    single_hash_level = action.get("single_hash_level", 0)  # Default: # = Title
    double_hash_level = action.get("double_hash_level", 1)  # Default: ## = Heading 1
    
    converted_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check for ## first (more specific)
        if text.startswith("## "):
            # Remove the ## and convert to heading
            new_text = text[3:].strip()
            para.text = new_text
            
            # Apply heading style
            if double_hash_level == 0:
                para.style = doc.styles['Title']
            else:
                para.style = doc.styles[f'Heading {double_hash_level}']
            
            converted_count += 1
            print(f"Converted '## {new_text}' to Heading {double_hash_level}")
            
        # Check for single # (but not ##)
        elif text.startswith("# ") and not text.startswith("## "):
            # Remove the # and convert to heading
            new_text = text[2:].strip()
            para.text = new_text
            
            # Apply heading style
            if single_hash_level == 0:
                para.style = doc.styles['Title']
            else:
                para.style = doc.styles[f'Heading {single_hash_level}']
            
            converted_count += 1
            print(f"Converted '# {new_text}' to Heading {single_hash_level}")
    
    print(f"Converted {converted_count} markdown headings to Word headings")

def apply_text_to_bullets_action(doc, elements_details, action):
    """
    Converts text enclosed in ** ** to bullet points
    Action: {"action": "convert_text_to_bullets", "bullet_style": "List Bullet", "remove_asterisks": true}
    """
    print(f"Applying text to bullets conversion: {action}")
    
    bullet_style = action.get("bullet_style", "List Bullet")
    remove_asterisks = action.get("remove_asterisks", True)
    
    converted_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check if text is enclosed in ** **
        if text.startswith("**") and text.endswith("**") and len(text) > 4:
            if remove_asterisks:
                # Remove the ** from both ends
                new_text = text[2:-2].strip()
                para.text = new_text
            
            # Apply bullet style
            try:
                para.style = doc.styles[bullet_style]
                converted_count += 1
                print(f"Converted '**{new_text}**' to bullet point")
            except KeyError:
                # If bullet style doesn't exist, create a simple bullet
                para.text = f"• {new_text}" if remove_asterisks else para.text
                converted_count += 1
                print(f"Converted to simple bullet (style '{bullet_style}' not found)")
    
    print(f"Converted {converted_count} items to bullet points")

def apply_pattern_to_style_action(doc, elements_details, action):
    """
    Generic pattern-to-style converter
    Action: {
        "action": "convert_pattern_to_style", 
        "pattern": "^# ", 
        "replacement": "", 
        "target_style": "Title",
        "pattern_type": "regex"
    }
    """
    print(f"Applying pattern to style conversion: {action}")
    
    import re
    
    pattern = action.get("pattern", "")
    replacement = action.get("replacement", "")
    target_style = action.get("target_style", "Normal")
    pattern_type = action.get("pattern_type", "regex")  # "regex" or "simple"
    
    if not pattern:
        print("Warning: No pattern specified for pattern_to_style conversion")
        return
    
    converted_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        
        if pattern_type == "regex":
            if re.match(pattern, text):
                # Apply replacement
                new_text = re.sub(pattern, replacement, text)
                para.text = new_text
                
                # Apply style
                try:
                    para.style = doc.styles[target_style]
                    converted_count += 1
                    print(f"Converted '{text}' to style '{target_style}'")
                except KeyError:
                    print(f"Warning: Style '{target_style}' not found")
        else:
            # Simple string matching
            if text.startswith(pattern):
                # Apply replacement
                new_text = text.replace(pattern, replacement, 1)
                para.text = new_text
                
                # Apply style
                try:
                    para.style = doc.styles[target_style]
                    converted_count += 1
                    print(f"Converted '{text}' to style '{target_style}'")
                except KeyError:
                    print(f"Warning: Style '{target_style}' not found")
    
    print(f"Converted {converted_count} items using pattern matching")