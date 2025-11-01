"""
Template-based document formatting utilities
"""
import os
import json
from docx import Document
from typing import Dict, Any
import re

def apply_business_proposal_template(input_doc_path: str, output_doc_path: str, template_path: str = None) -> Dict[str, Any]:
    """
    Apply business proposal template to unstructured content
    
    Args:
        input_doc_path: Path to input document with raw content
        output_doc_path: Path where formatted proposal should be saved
        template_path: Path to template (optional, uses default if not provided)
    
    Returns:
        Dict with operation results
    """
    try:
        # Default template path
        if not template_path:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(current_dir, "..", "templates", "business_proposal_template.docx")
        
        # Check if template exists
        if not os.path.exists(template_path):
            return {
                "success": False,
                "error": f"Template not found at {template_path}",
                "suggestion": "Please place your business proposal template at the specified location"
            }
        
        # Load input document
        input_doc = Document(input_doc_path)
        input_text = ""
        for paragraph in input_doc.paragraphs:
            input_text += paragraph.text + "\n"
        
        # Extract content using simple keyword matching
        # This is a basic implementation - you can enhance with AI extraction
        extracted_content = extract_business_content(input_text)
        
        # Load template
        template_doc = Document(template_path)
        
        # Replace placeholders in template
        replace_placeholders_in_doc(template_doc, extracted_content)
        
        # Save the result
        template_doc.save(output_doc_path)
        
        return {
            "success": True,
            "input_doc_path": input_doc_path,
            "output_doc_path": output_doc_path,
            "template_used": template_path,
            "placeholders_replaced": list(extracted_content.keys()),
            "message": "Business proposal template applied successfully"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "input_doc_path": input_doc_path,
            "output_doc_path": output_doc_path
        }

def extract_business_content(text: str) -> Dict[str, str]:
    """
    Extract business proposal content from unstructured text
    This is a basic implementation - can be enhanced with AI
    """
    content = {}
    
    # Simple keyword-based extraction
    # You can enhance this with AI-powered extraction
    
    # Company name (look for patterns)
    company_match = re.search(r'(?:company|organization|firm):\s*([^\n]+)', text, re.IGNORECASE)
    if company_match:
        content['{{COMPANY_NAME}}'] = company_match.group(1).strip()
    
    # Project title
    title_match = re.search(r'(?:project|proposal|title):\s*([^\n]+)', text, re.IGNORECASE)
    if title_match:
        content['{{PROJECT_TITLE}}'] = title_match.group(1).strip()
    
    # Executive summary
    summary_match = re.search(r'(?:executive summary|summary|overview):\s*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if summary_match:
        content['{{EXECUTIVE_SUMMARY}}'] = summary_match.group(1).strip()
    
    # Problem statement
    problem_match = re.search(r'(?:problem|challenge|issue):\s*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if problem_match:
        content['{{PROBLEM_STATEMENT}}'] = problem_match.group(1).strip()
    
    # Solution
    solution_match = re.search(r'(?:solution|approach|methodology):\s*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if solution_match:
        content['{{PROPOSED_SOLUTION}}'] = solution_match.group(1).strip()
    
    # Timeline
    timeline_match = re.search(r'(?:timeline|schedule|duration):\s*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if timeline_match:
        content['{{TIMELINE}}'] = timeline_match.group(1).strip()
    
    # Budget
    budget_match = re.search(r'(?:budget|cost|price|investment):\s*([^\n]+(?:\n[^\n]+)*)', text, re.IGNORECASE)
    if budget_match:
        content['{{BUDGET}}'] = budget_match.group(1).strip()
    
    # Default values for missing content
    defaults = {
        '{{COMPANY_NAME}}': '[Company Name]',
        '{{PROJECT_TITLE}}': '[Project Title]',
        '{{EXECUTIVE_SUMMARY}}': '[Executive Summary]',
        '{{PROBLEM_STATEMENT}}': '[Problem Statement]',
        '{{PROPOSED_SOLUTION}}': '[Proposed Solution]',
        '{{TIMELINE}}': '[Timeline]',
        '{{BUDGET}}': '[Budget Information]',
        '{{TEAM_MEMBERS}}': '[Team Information]',
        '{{CONTACT_INFO}}': '[Contact Information]'
    }
    
    # Fill in defaults for missing content
    for placeholder, default_value in defaults.items():
        if placeholder not in content:
            content[placeholder] = default_value
    
    return content

def replace_placeholders_in_doc(doc: Document, replacements: Dict[str, str]):
    """
    Replace placeholders in a Word document
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)
    
    # Replace in headers/footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)